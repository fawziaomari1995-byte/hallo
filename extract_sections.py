"""Extract structured headings and section text from the generated DOCX report.
This helper script is useful for quickly reviewing the report outline and verifying
that key sections are present in the Word document.
"""

from docx import Document

DOC = 'التقرير-1.docx'
try:
    doc = Document(DOC)
except Exception as e:
    print('ERROR_OPEN', e)
    raise SystemExit(1)

# Identify heading paragraphs by style name starting with 'Heading' or 'Title'
headings = []
for i, p in enumerate(doc.paragraphs):
    style_name = getattr(p.style, 'name', '')
    if style_name and (style_name.startswith('Heading') or style_name == 'Title' or style_name == 'Intense Quote'):
        headings.append((i, p.text.strip(), style_name))

# If no headings found by style, fall back to known Arabic section titles
if not headings:
    candidates = ['تصميم قاعدة البيانات','تصميم الواجهات','مخطط حالات الاستخدام','الخوارزميات والعمليات','الاختبارات','التقنيات المستخدمة','اللغات المستخدمة','الاستنتاجات والتوصيات','ملحق الملفات','المقدمة','الإهداء']
    for i,p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt in candidates:
            headings.append((i, txt, 'Detected'))

print('FOUND HEADINGS:', len(headings))

sections = []
for idx, title, style in headings:
    # collect paragraphs after idx until next heading
    content_lines = []
    j = idx + 1
    while j < len(doc.paragraphs):
        sname = getattr(doc.paragraphs[j].style, 'name', '')
        if sname and (sname.startswith('Heading') or sname == 'Title' or sname == 'Intense Quote'):
            break
        # also break if paragraph text is exactly another known heading
        txt = doc.paragraphs[j].text.strip()
        if txt and txt in [h[1] for h in headings]:
            break
        content_lines.append(txt)
        j += 1
    sections.append((title, '\n'.join(content_lines).strip()))

# Print sections requested: show full text for each
for title, body in sections:
    print('\n' + '='*40)
    print('SECTION:', title)
    print('-'*40)
    if body:
        print(body)
    else:
        print('[بدون محتوى نصي مباشر تحت هذا العنوان]')

# If the specific 'مخطط حالات الاستخدام' section exists, print it separately
for title, body in sections:
    if 'مخطط حالات الاستخدام' in title:
        print('\n-- Use case section found and printed above.')
        break
else:
    print('\n-- Use case section not found by heading styles.')
