"""Script to add code explanations to the DOCX file"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Open the existing document
doc = Document('التقرير-1-UPDATED.docx')

# Add the new section at the end
doc.add_paragraph()  # Add spacing

# Add the main heading
heading = doc.add_heading('شرح الأكواس البرمجية', level=1)
heading_format = heading.paragraph_format
heading_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# Introduction
intro = doc.add_paragraph('يقدم هذا القسم شرحاً تفصيلياً للأكواد البرمجية الأساسية المستخدمة في نظام إدارة الفعاليات.')
intro.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# ========== 1. Server.js Section ==========
doc.add_heading('1. ملف الخادم الرئيسي (server.js)', level=2)

p = doc.add_paragraph('الوظيفة الأساسية:')
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
p = doc.add_paragraph('ملف الخادم الذي يدير جميع طلبات المستخدمين والعمليات الخلفية للنظام باستخدام Express.js')
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

p = doc.add_paragraph('المكتبات المستخدمة:', style='List Bullet')
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

libraries = [
    'express: إطار العمل الويب لإنشاء REST APIs',
    'sqlite3: قاعدة البيانات المحلية',
    'bcrypt: تشفير كلمات المرور',
    'multer: رفع الملفات والصور',
    'nodemailer: إرسال الرسائل البريدية',
    'QRCode: توليد رموز QR للتذاكر',
    'PDFKit: توليد ملفات PDF',
    'Stripe: معالجة الدفع الآمنة'
]

for lib in libraries:
    p = doc.add_paragraph(lib, style='List Bullet')
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# خطوات الإعداد
doc.add_heading('خطوات الإعداد الأساسية:', level=3)

steps = [
    ('إعداد التطبيق والمنافذ', 'const PORT = process.env.PORT || 3000; يتم تعيين منفذ الخادم (3000 افتراضياً)'),
    ('إعداد مجلد قاعدة البيانات', 'في بيئة الإنتاج (Render) يتم استخدام /tmp، وفي التطوير المحلي يتم استخدام مجلد db'),
    ('تخزين الملفات المرفوعة', 'استخدام multer لإنشاء أسماء ملفات فريدة وتجنب التضاربات'),
    ('المستخدمون الافتراضيون', 'admin/admin123 و user/user123 للاختبار')
]

for step_title, step_desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run(step_title + ': ')
    run.bold = True
    p.add_run(step_desc)

# ========== 2. db.js Section ==========
doc.add_heading('2. ملف قاعدة البيانات (db.js)', level=2)

p = doc.add_paragraph('الوظيفة: إعداد واتصال قاعدة بيانات SQLite المشتركة عبر التطبيق')
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

points = [
    'إنشاء مجلد db إذا لم يكن موجوداً',
    'فتح الاتصال مع قاعدة البيانات events.db',
    'إيقاف التطبيق فوراً عند فشل الاتصال',
    'الاتصال يبقى مفتوحاً طوال حياة التطبيق'
]

for point in points:
    p = doc.add_paragraph(point, style='List Bullet')
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# ========== 3. server-utils.js Section ==========
doc.add_heading('3. ملف المساعدات - المصادقة والأمان (lib/server-utils.js)', level=2)

p = doc.add_paragraph('الوظائف الأساسية:')
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

utils_functions = [
    ('generateToken()', 'توليد رمز عشوائي فريد للجلسات والمصادقة'),
    ('getResetPasswordUrl()', 'بناء رابط إعادة تعيين كلمة المرور للبريد الإلكتروني'),
    ('getActivationUrl()', 'بناء رابط تفعيل الحساب'),
    ('buildActivationEmailHtml()', 'إنشاء محتوى رسالة البريد الترحيبية بصيغة HTML'),
    ('makeAttachOptionalUser()', 'Middleware اختياري: يضيف معلومات المستخدم إن وجدت'),
    ('makeRequireAuth()', 'Middleware إجباري: يرفض الطلب إذا لم يكن المستخدم مصرح'),
    ('makeRequireAdmin()', 'Middleware إجباري: يسمح فقط للمديرين')
]

for func_name, func_desc in utils_functions:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run(func_name + ': ')
    run.bold = True
    p.add_run(func_desc)

# ========== 4. pdf-utils.js Section ==========
doc.add_heading('4. ملف توليد التذاكر (lib/pdf-utils.js)', level=2)

p = doc.add_paragraph('الوظيفة: توليد ملفات PDF للتذاكر تتضمن رمز QR للتحقق عند الدخول')
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

doc.add_heading('خطوات العملية:', level=3)

pdf_steps = [
    'إنشاء مستند PDF بحجم A6 (بطاقة صغيرة)',
    'إضافة معلومات الفعالية: العنوان والنوع والمقعد',
    'تفريق بين التذاكر الافتراضية (virtual) والحضورية',
    'إضافة رمز QR في وسط البطاقة',
    'حفظ الملف في النظام وإرساله للمستخدم'
]

for step in pdf_steps:
    p = doc.add_paragraph(step, style='List Number')
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# ========== 5. العمليات الأساسية ==========
doc.add_heading('5. العمليات الأساسية', level=2)

doc.add_heading('عملية الحجز والدفع:', level=3)

booking_steps = [
    'المستخدم يختار عدد التذاكر والنوع المطلوب',
    'التحقق من التوافر: هل عدد التذاكر متوفر؟',
    'حجز مقعد مؤقت لتفادي البيع المتزامن',
    'إنشاء جلسة دفع آمنة عبر Stripe',
    'عرض نموذج الدفع والمستخدم يدخل بيانات البطاقة',
    'معالجة الدفع والتحقق من النتيجة عبر Webhooks',
    'إذا نجح: حفظ الحجز وتوليد رمز QR فريد',
    'توليد ملف PDF للتذكرة وإرساله في البريد الإلكتروني'
]

for step in booking_steps:
    p = doc.add_paragraph(step, style='List Number')
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

doc.add_heading('عملية المصادقة والتسجيل:', level=3)

auth_steps = [
    'المستخدم يدخل بيانات التسجيل (بريد وكلمة مرور)',
    'تشفير كلمة المرور باستخدام bcrypt',
    'حفظ البيانات في قاعدة البيانات',
    'إرسال رسالة تفعيل عبر البريد الإلكتروني',
    'المستخدم يضغط على رابط التفعيل',
    'توليد توكن (رمز جلسة) فريد',
    'المستخدم يرسل التوكن في رؤوس كل طلب (Authorization: Bearer token)'
]

for step in auth_steps:
    p = doc.add_paragraph(step, style='List Number')
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# ========== 6. معايير الأمان ==========
doc.add_heading('6. معايير الأمان المطبقة', level=2)

security_points = [
    ('تشفير كلمات المرور', 'استخدام bcrypt مع 10 جولات تشفير'),
    ('المصادقة بالتوكن', 'كل طلب يحتوي على Bearer Token يتم التحقق من صحته'),
    ('الصلاحيات المختلفة', 'Admin للإدارة، Normal User للحجز والدفع فقط'),
    ('معالجة الدفع الآمنة', 'استخدام Stripe Elements (بيانات البطاقة لا ترسل للخادم)'),
    ('التحقق عبر Webhooks', 'Stripe يرسل تأكيد الدفع مباشرة للخادم'),
    ('صلاحية الرموز المؤقتة', 'روابط إعادة تعيين كلمة المرور صالحة لساعة واحدة فقط')
]

for sec_title, sec_desc in security_points:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run(sec_title + ': ')
    run.bold = True
    p.add_run(sec_desc)

# ========== 7. المتغيرات البيئية ==========
doc.add_heading('7. المتغيرات البيئية (Environment Variables)', level=2)

env_vars = [
    ('PORT', 'منفذ الخادم (افتراضي: 3000)'),
    ('APP_URL', 'رابط التطبيق الكامل (مثل https://example.com)'),
    ('STRIPE_SECRET_KEY', 'مفتاح Stripe السري للدفع'),
    ('STRIPE_PUBLISHABLE_KEY', 'مفتاح Stripe العام للعميل'),
    ('SMTP_HOST', 'خادم البريد الإلكتروني'),
    ('SMTP_USER', 'اسم مستخدم البريد'),
    ('SMTP_PASS', 'كلمة مرور البريد'),
    ('SMTP_PORT', 'منفذ البريد (افتراضي: 587)'),
    ('RENDER', 'تحديد بيئة الاستضافة (Render vs محلي)')
]

for var_name, var_desc in env_vars:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run(var_name + ': ')
    run.bold = True
    p.add_run(var_desc)

# ========== 8. نصائح للتطوير ==========
doc.add_heading('8. نصائح للتطوير والاختبار', level=2)

doc.add_heading('تشغيل الخادم محلياً:', level=3)
p = doc.add_paragraph()
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
run = p.add_run('npm install')
run.italic = True
p.add_run('  لتثبيت المكتبات')

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
run = p.add_run('npm start')
run.italic = True
p.add_run('  لتشغيل الخادم على المنفذ 3000')

doc.add_heading('اختبار المصادقة:', level=3)
tips = [
    'طلب بدون مصادقة سيرجع خطأ 401 Unauthorized',
    'طلب بـ Bearer Token صحيح سيعيد البيانات المطلوبة',
    'استخدم أدوات مثل Postman أو curl لاختبار الـ APIs'
]
for tip in tips:
    p = doc.add_paragraph(tip, style='List Bullet')
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

doc.add_heading('اختبار الدفع:', level=3)
payment_tips = [
    'استخدم مفاتيح Stripe للبيئة الاختبارية (Test Keys)',
    'استخدم أرقام بطاقات الاختبار من Stripe (مثل 4242 4242 4242 4242)',
    'لا تستخدم بيانات حقيقية في بيئة الاختبار'
]
for tip in payment_tips:
    p = doc.add_paragraph(tip, style='List Bullet')
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# ========== 9. استكشاف الأخطاء ==========
doc.add_heading('9. استكشاف الأخطاء الشائعة', level=2)

errors = [
    ('Unable to open database', 'تحقق من صلاحيات مجلد db أو المساحة المتاحة على القرص'),
    ('Unauthorized (401)', 'تحقق من وجود التوكن في رؤوس الطلب أو انتهاء صلاحيته'),
    ('Forbidden (403)', 'المستخدم الحالي ليس لديه صلاحيات كافية (مثل محاولة عملية تحتاج صلاحية مدير)'),
    ('Payment failed', 'تحقق من مفاتيح Stripe أو الاتصال بالإنترنت أو صحة بيانات البطاقة'),
    ('Email not sent', 'تحقق من إعدادات SMTP وتأكد من صحة البيانات')
]

for error_name, solution in errors:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run(error_name + ': ')
    run.bold = True
    p.add_run(solution)

# Add footer
doc.add_paragraph()
p = doc.add_paragraph('آخر تحديث: 2026-06-20')
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Save the document
doc.save('التقرير-1-UPDATED.docx')
print('✅ تم إضافة شروح الأكواد بنجاح إلى الملف!')
print('📄 الملف: التقرير-1-UPDATED.docx')
