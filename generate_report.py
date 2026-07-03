# -*- coding: utf-8 -*-
import os
import re
import zipfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
try:
    import arabic_reshaper
    from bidi.algorithm import get_display
except Exception:
    arabic_reshaper = None
    get_display = None
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
import shutil
import json

"""Generate an Arabic project report in DOCX and PDF formats.
This script collects project metadata, placeholder diagrams, and structured sections
into a formatted report suitable for the Daraa event management project.
"""

WORKSPACE_ROOT = os.path.dirname(__file__)
# Use fixed output filenames (user requested ثابت names). Remove existing files before writing to avoid PermissionError.
OUTPUT_DOCX = os.path.join(WORKSPACE_ROOT, 'التقرير-1.docx')
OUTPUT_PDF = os.path.join(WORKSPACE_ROOT, 'التقرير-1.pdf')
OUTPUT_ZIP = os.path.join(WORKSPACE_ROOT, 'التقرير-1.zip')
IMAGES_DIR = os.path.join(WORKSPACE_ROOT, 'images')

PROJECT_TITLE = 'موقع فعاليات محافظة درعا'
ACADEMIC_YEAR = '2024 - 2025'
INSTITUTE_TITLE = 'المعهد التقاني للحاسوب بدرعا'
DEPARTMENT_TITLE = 'قسم هندسة البرمجيات'
TEAM_MEMBERS = [
    'محمد يزن عياش',
    'شيماء محمد',
    'شفاء الشليب',
    'ديانا الكور',
    'غزل البشري',
    'هبا الفرا',
    'سدرة إبراهيم'
]

ARABIC_INTRO = (
    'يقدم هذا التقرير وصفاً تفصيلياً لمشروع الموقع الإلكتروني لإدارة فعاليات محافظة درعا، ' 
    'ويغطي الهيكل العام للنظام، المكونات التقنية، التصميم، قاعدة البيانات، واجهات المستخدم، ' 
    'والاختبارات المنفذة. يركز التقرير على إبراز قدرات النظام في الحجز الإلكتروني، إدارة التذاكر، ' 
    'التواصل مع المستخدمين، وأتمتة العمليات الإدارية.'
)


# Find a usable Arabic TrueType font on the local machine or in the project.
# This font is used when generating images and PDF content in Arabic.
def find_arabic_font():
    # 1) check project-local fonts folder or any .ttf in workspace root
    fonts_dirs = [os.path.join(WORKSPACE_ROOT, 'fonts'), WORKSPACE_ROOT]
    for fd in fonts_dirs:
        if os.path.exists(fd):
            for f in os.listdir(fd):
                if f.lower().endswith('.ttf'):
                    return os.path.join(fd, f)
    # 2) fallback to common Windows-installed Arabic fonts
    candidates = [
        r"C:\\Windows\\Fonts\\TraditionalArabic.ttf",
        r"C:\\Windows\\Fonts\\Scheherazade-Regular.ttf",
        r"C:\\Windows\\Fonts\\Amiri-Regular.ttf",
        r"C:\\Windows\\Fonts\\NotoNaskhArabic-Regular.ttf",
        r"C:\\Windows\\Fonts\\arial.ttf",
        r"C:\\Windows\\Fonts\\tahoma.ttf",
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return None

ARABIC_FONT_PATH = find_arabic_font()


# Transform Arabic text for correct right-to-left rendering when drawing images.
def shape_arabic(text):
    """Reshape and apply bidi to Arabic text for proper display in PIL images."""
    if not text:
        return text
    if arabic_reshaper and get_display:
        try:
            reshaped = arabic_reshaper.reshape(text)
            bidi_text = get_display(reshaped)
            return bidi_text
        except Exception:
            return text
    return text


# Read package.json to include project metadata in the report.
def read_package_json():
    import json
    path = os.path.join(WORKSPACE_ROOT, 'package.json')
    if not os.path.exists(path):
        return {}
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return {}


# Summarize the content and basic stats of a file for inclusion in the report appendix.
def summarize_file(path):
    try:
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception:
        try:
            return {'lines': 0, 'size': os.path.getsize(path), 'preview': '', 'summary': 'ملف ثنائي أو غير قابل للقراءة كنص.'}
        except OSError:
            return {'lines': 0, 'size': 0, 'preview': '', 'summary': 'ملف غير قابل للقراءة.'}
    lines = content.splitlines()
    preview = '\n'.join(lines[:40])
    funcs = re.findall(r'def\s+(\w+)|function\s+(\w+)|\b(\w+)\s*=\s*\(.*?\)\s*=>', content)
    func_names = set()
    for t in funcs:
        for name in t:
            if name:
                func_names.add(name)
    summary = [f"عدد الأسطر: {len(lines)}"]
    if func_names:
        summary.append('الدوال/المكونات الرئيسية: ' + ', '.join(list(func_names)[:10]))
    if path.endswith('.html'):
        tags = len(re.findall(r'<\s*[^>]+>', content))
        summary.append(f"عناصر HTML تقريبية: {tags}")
    if path.endswith('.js'):
        vars_found = len(re.findall(r'var\s+|let\s+|const\s+', content))
        summary.append(f"تعريفات متغيرات تقريبية: {vars_found}")
    if path.endswith('.css'):
        selectors = len(re.findall(r'\w+\s*{', content))
        summary.append(f"محددات CSS تقريبية: {selectors}")
    return {'lines': len(lines), 'size': os.path.getsize(path), 'preview': preview, 'summary': ' - '.join(summary)}


# Collect a sorted list of workspace files to include in the project appendix.
def collect_files(root):
    files = []
    for dirpath, dirnames, filenames in os.walk(root):
        skip = ['.git', '__pycache__', 'node_modules', '.vscode']
        if any(p in dirpath for p in skip):
            continue
        for fname in filenames:
            full = os.path.join(dirpath, fname)
            rel = os.path.relpath(full, root)
            files.append((full, rel))
    files.sort(key=lambda x: x[1])
    return files


# Generate a fallback ERD image when a real diagram is not available.
def create_sample_erd(path):
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return False
    img = Image.new('RGB', (1200, 900), color='white')
    draw = ImageDraw.Draw(img)
    boxes = [
        ((120, 120), (420, 260), 'Users'),
        ((520, 120), (820, 260), 'Events'),
        ((320, 340), (620, 480), 'Bookings'),
        ((720, 340), (1020, 480), 'Tickets'),
    ]
    for b in boxes:
        draw.rectangle([b[0], b[1]], outline='black', width=3)
        try:
            if ARABIC_FONT_PATH:
                font = ImageFont.truetype(ARABIC_FONT_PATH, 24)
            else:
                font = ImageFont.truetype('arial.ttf', 24)
        except Exception:
            font = ImageFont.load_default()
        text_to_draw = shape_arabic(b[2])
        bbox = draw.textbbox((0, 0), text_to_draw, font=font)
        w = bbox[2] - bbox[0]
        h = bbox[3] - bbox[1]
        cx = (b[0][0] + b[1][0]) // 2 - w // 2
        cy = (b[0][1] + b[1][1]) // 2 - h // 2
        draw.text((cx, cy), text_to_draw, fill='black', font=font)
    draw.line((420, 190, 520, 190), fill='black', width=3)
    draw.polygon([(520, 190), (510, 185), (510, 195)], fill='black')
    draw.line((620, 415, 720, 415), fill='black', width=3)
    draw.polygon([(720, 415), (710, 410), (710, 420)], fill='black')
    img.save(path)
    return True


# Create a simple use-case diagram image to illustrate system actors and actions.
def create_use_case_diagram(path):
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return False
    img = Image.new('RGB', (1200, 900), color='white')
    draw = ImageDraw.Draw(img)
    try:
        if ARABIC_FONT_PATH:
            font = ImageFont.truetype(ARABIC_FONT_PATH, 18)
        else:
            font = ImageFont.truetype('arial.ttf', 18)
    except Exception:
        font = ImageFont.load_default()
    # Actors
    draw.text((80, 120), 'المستخدم', fill='black', font=font)
    draw.ellipse((40, 120, 100, 180), outline='black', width=2)
    draw.text((80, 320), 'المنظم', fill='black', font=font)
    draw.ellipse((40, 320, 100, 380), outline='black', width=2)
    # Use cases
    ovals = [
        ((400, 120), (760, 180), 'عرض الفعاليات'),
        ((400, 220), (760, 280), 'حجز تذكرة'),
        ((400, 320), (760, 380), 'دفع عبر Stripe'),
        ((400, 420), (760, 480), 'توليد تذكرة QR'),
        ((400, 520), (760, 580), 'إدارة الحساب'),
    ]
    for b in ovals:
        draw.ellipse([b[0], b[1]], outline='black', width=2)
        text_to_draw = shape_arabic(b[2])
        bbox = draw.textbbox((0, 0), text_to_draw, font=font)
        w = bbox[2] - bbox[0]
        h = bbox[3] - bbox[1]
        cx = (b[0][0] + b[1][0]) // 2 - w // 2
        cy = (b[0][1] + b[1][1]) // 2 - h // 2
        draw.text((cx, cy), text_to_draw, fill='black', font=font)
    connections = [
        ((100, 150), (400, 150)),
        ((100, 350), (400, 350)),
    ]
    for line in connections:
        draw.line(line, fill='black', width=2)
    img.save(path)
    return True


# Generate placeholder UI screenshot images for report visuals.
def create_ui_screenshot(path, title, subtext):
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return False
    img = Image.new('RGB', (1200, 700), color='#f5f5f5')
    draw = ImageDraw.Draw(img)
    try:
        if ARABIC_FONT_PATH:
            font = ImageFont.truetype(ARABIC_FONT_PATH, 28)
            small = ImageFont.truetype(ARABIC_FONT_PATH, 18)
        else:
            font = ImageFont.truetype('arial.ttf', 28)
            small = ImageFont.truetype('arial.ttf', 18)
    except Exception:
        font = ImageFont.load_default()
        small = font
    draw.rectangle((50, 50, 1150, 650), fill='white', outline='black', width=2)
    draw.text((70, 70), shape_arabic(title), fill='black', font=font)
    draw.rectangle((70, 130, 1130, 210), fill='#e8f0ff', outline='black', width=1)
    draw.text((90, 150), shape_arabic(subtext), fill='black', font=small)
    sections = ['شريط تنقل', 'قائمة فعاليات', 'تفاصيل مختصرة', 'زر حجز']
    y = 240
    for sec in sections:
        draw.rectangle((90, y, 1110, y + 70), outline='#d1d1d1', width=1)
        draw.text((110, y + 24), shape_arabic(sec), fill='black', font=small)
        y += 90
    img.save(path)
    return True


# Build a set of placeholder UI images used by the report when real screenshots are absent.
def create_ui_images():
    items = [
        ('ui_index.png', 'الصفحة الرئيسية', 'عرض الفعاليات المتاحة مع فلتر الموقع والتصنيف.'),
        ('ui_event.png', 'صفحة تفاصيل الفعالية', 'معلومات كاملة عن الفعالية، المواعيد، والموقع.'),
        ('ui_checkout.png', 'صفحة الدفع', 'نافذة إدخال تفاصيل الدفع وخيارات التذاكر.'),
        ('ui_profile.png', 'صفحة حساب المستخدم', 'عرض بيانات المستخدم، الحجوزات، والتاريخ.'),
    ]
    created = []
    for filename, title, subtext in items:
        path = os.path.join(IMAGES_DIR, filename)
        if create_ui_screenshot(path, title, subtext):
            created.append(path)
    return created


# Copy real screenshot files into the report image directory if they exist.
def collect_real_screenshots():
    """If user placed real screenshots in `screenshots/`, copy them into `images/`.
    Accepts png/jpg and returns list of copied image paths. If none found, returns []."""
    src_dir = os.path.join(WORKSPACE_ROOT, 'screenshots')
    if not os.path.exists(src_dir):
        return []
    copied = []
    for fname in os.listdir(src_dir):
        if not fname.lower().endswith(('.png', '.jpg', '.jpeg')):
            continue
        src = os.path.join(src_dir, fname)
        dst = os.path.join(IMAGES_DIR, fname)
        try:
            shutil.copy2(src, dst)
            copied.append(dst)
        except Exception:
            continue
    return copied


# Load optional custom section content from a JSON override file.
def load_section_overrides():
    """Load `section_overrides.json` if exists. Returns dict of overrides."""
    path = os.path.join(WORKSPACE_ROOT, 'section_overrides.json')
    if not os.path.exists(path):
        return {}
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return {}


# Apply consistent formatting and colors to tables inside the Word report.
def style_table(table, col_widths=None, header_fill='D9D9D9'):
    """Apply a consistent style to a python-docx table: Table Grid, optional column widths, header shading, and alternate row styling."""
    try:
        table.style = 'Table Grid'
    except Exception:
        pass
    try:
        table.autofit = False
    except Exception:
        pass
    # set column widths if provided
    if col_widths:
        from docx.shared import Inches
        for i, w in enumerate(col_widths):
            try:
                for row in table.rows:
                    row.cells[i].width = Inches(w)
            except Exception:
                continue
    # shade header row and align text
    try:
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), header_fill)
            tcPr.append(shd)
    except Exception:
        pass
    # apply alternate row shading and right alignment for body rows
    try:
        for row_index, row in enumerate(table.rows[1:], start=1):
            if row_index % 2 == 0:
                fill = 'F2F2F2'
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), fill)
                    tcPr.append(shd)
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    except Exception:
        pass


# Add the report cover page with title, institute name, and student team members.
def add_cover_page(doc):
    doc.add_paragraph('1', style='Title')
    doc.add_paragraph('وزارة التعليم العالي', style='Title')
    doc.add_paragraph('جامعة دمشق', style='Title')
    doc.add_paragraph(INSTITUTE_TITLE, style='Title')
    doc.add_paragraph(DEPARTMENT_TITLE, style='Title')
    doc.add_paragraph('\n')
    doc.add_paragraph(PROJECT_TITLE, style='Title')
    doc.add_paragraph('\n')
    doc.add_paragraph('أعدّه الطلاب التالية أسماؤهم:', style='Intense Quote')
    for name in TEAM_MEMBERS:
        doc.add_paragraph(f'- {name}', style='List Bullet')
    doc.add_paragraph('\n')
    doc.add_paragraph(f'العام الدراسي {ACADEMIC_YEAR}', style='Intense Quote')
    doc.add_section(WD_SECTION.NEW_PAGE)


# Create the table of contents page with section titles and page placeholders.
def add_contents_page(doc):
    doc.add_heading('فهرس المحتويات', level=1)
    entries = [
        ('الإهداء', 3),
        ('المقدمة', 4),
        ('الهدف من المشروع', 4),
        ('مزايا المشروع', 5),
        ('التقنيات المستخدمة', 6),
        ('شرح الكود', 7),
        ('اللغات المستخدمة', 8),
        ('تصميم قاعدة البيانات', 9),
        ('تصميم الواجهات', 12),
        ('الخوارزميات والعمليات', 14),
        ('الاختبارات', 16),
        ('الاستنتاجات والتوصيات', 18),
        ('ملحق الملفات', 20),
    ]
    for title, page in entries:
        p = doc.add_paragraph(style='Body Text')
        p.add_run(title)
        p.add_run('.').bold = False
        p.add_run(' ' * 20)
        p.add_run(str(page)).bold = False
    doc.add_section(WD_SECTION.NEW_PAGE)


# Add the dedication section to the report.
def add_dedication(doc):
    doc.add_heading('الإهداء', level=1)
    doc.add_paragraph('"وأخير دعواهم أن الحمد لله رب العالمين"')
    doc.add_paragraph('إلى كل من ساندنا ووقف معنا في هذه الرحلة العلمية الطويلة، إلى أمهاتنا وآبائنا، وإلى أساتذتنا الكرام. شكراً لكل كلمة تشجيع ولكل لحظة دعم، ولكل من آمن بنا ووقف معنا حتى انتهى المشروع.')
    doc.add_section(WD_SECTION.NEW_PAGE)


# Add the introduction section, with optional override support from JSON.
def add_introduction(doc):
    doc.add_heading('المقدمة', level=1)
    overrides = load_section_overrides()
    intro_text = overrides.get('المقدمة', ARABIC_INTRO)
    doc.add_paragraph(intro_text)
    doc.add_section(WD_SECTION.NEW_PAGE)


# Add project objectives as a numbered list section.
def add_objectives(doc):
    doc.add_heading('الهدف من المشروع', level=1)
    items = [
        'إنشاء نظام إلكتروني لإدارة فعاليات محافظة درعا بصورة متكاملة.',
        'تيسير حجز التذاكر ومتابعة الحجوزات إلكترونياً.',
        'إدارة بيانات الفعاليات والمستخدمين بأمان وفعالية.',
        'توليد تذاكر QR وتقارير مبيعات تلقائية.',
        'تمكين الدفع عبر الإنترنت باستخدام Stripe.',
    ]
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.add_run(item)
    doc.add_section(WD_SECTION.NEW_PAGE)


# Add a section that explains the advantages of the system.
def add_advantages(doc):
    doc.add_heading('مزايا المشروع', level=1)
    advantages = [
        'زيادة الكفاءة وتقليل الاعتماد على الإجراءات اليدوية.',
        'تحسين تجربة المستخدم في متابعة فعاليات درعا وحجز التذاكر.',
        'توفير معلومات دقيقة وفورية عن توافر المقاعد والإيرادات.',
        'أمان أكبر للبيانات بفضل استخدام قاعدة بيانات محلية منظمة.',
        'سهولة توسعة النظام وإضافة خصائص جديدة مستقبلاً.',
    ]
    for item in advantages:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run('🌷 ' + item)
    doc.add_section(WD_SECTION.NEW_PAGE)


# Add the technologies section describing the stack used in the project.
def add_technologies(doc, package_info):
    doc.add_heading('التقنيات المستخدمة', level=1)
    techs = [
        ('Node.js', 'بيئة تشغيل جافاسكربت على الخادم. لماذا: أداء جيد للمشروعات الخفيفة/المتوسطة، مجتمع واسع. البدائل: Python/Flask, Ruby on Rails. إعدادات مقترحة: استخدام `pm2` أو `systemd` للتشغيل كخدمة.'),
        ('Express', 'إطار عمل ويب خفيف لبناء REST APIs وواجهات الـ HTTP. لماذا: بسيط ومرن، ملائم لمشروع تعليمي. ممارسات: تنظيم المسارات، استخدام middlewares للتحقق والأمان.'),
        ('SQLite', 'قاعدة بيانات ملفية مناسبة لبيئات التطوير والنشر البسيط. لماذا:无需 خادم منفصل، سهولة النسخ الاحتياطي. اعتبارات: لا تناسب أحمال كتابة عالية؛ للبيئات الإنتاجية الكبيرة استخدم PostgreSQL.'),
        ('HTML / CSS / JavaScript', 'تقنيات الواجهة الأمامية لعرض المحتوى والتفاعل. ممارسات مقترحة: استخدام تصميم مستجيب (responsive) واعتبارات الوصول (accessibility).'),
        ('Stripe', 'مزود دفع آمن لمعالجة بطاقات الائتمان. الإعداد: مفاتيح API بيئة اختبار/إنتاج، webhook لمعالجة الأحداث (payment_intent.succeeded). بدائل: PayPal، local gateway.'),
        ('PDFKit / QRCode', 'مكتبات لتوليد ملفات PDF ورموز QR. الاستخدام: توليد تذاكر PDF مرفقة برمز QR فريد لحصول الدخول.'),
        ('Nodemailer', 'مكتبة إرسال البريد الإلكتروني. الإعداد: SMTP أو خدمات مثل SendGrid مع مفاتيح آمنة.'),
        ('Multer', 'لتحميل الملفات (صور الفعالية). ممارسات: تحديد حجم الملفات المسموح والتخزين المؤقت والتحقق من النوع MIME.'),
    ]
    for name, desc in techs:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(f'{name}: {desc}')
        try:
            if ARABIC_FONT_PATH:
                r.font.name = 'Traditional Arabic'
        except Exception:
            pass
    if package_info.get('description'):
        doc.add_paragraph('\nوصف المشروع من package.json:')
        doc.add_paragraph(package_info['description'], style='Body Text')
    # allow override content for this section
    overrides = load_section_overrides()
    if overrides.get('التقنيات المستخدمة'):
        doc.add_paragraph('\n')
        doc.add_paragraph(overrides['التقنيات المستخدمة'], style='Body Text')
    doc.add_section(WD_SECTION.NEW_PAGE)


# Add the code explanation section that summarizes front-end, back-end, and database choices.
def add_code_explanation(doc):
    doc.add_heading('شرح الكود', level=1)
    doc.add_paragraph('يشرح هذا القسم مكونات البرمجة الأساسية في المشروع ولماذا تم اختيارها.')
    doc.add_paragraph('الواجهة الأمامية (Front-end) تعتمد على HTML وCSS وJavaScript لتقديم تجربة مستخدم تفاعلية وسهلة الاستخدام. تم استخدام HTML لهيكلة الصفحات، وCSS لتهيئة التصميم، وJavaScript لإدارة الأحداث، تحديث المحتوى دون إعادة تحميل كامل، والتفاعل مع واجهات API.')
    doc.add_paragraph('الجهة الخلفية (Back-end) مبنية باستخدام Node.js مع Express. يوفر هذا المزيج قدرة جيدة على معالجة طلبات HTTP، التحقق من البيانات، إدارة جلسات المستخدم، والتكامل مع خدمات الدفع مثل Stripe. ميزة استخدام Node.js هي القدرة على استخدام نفس لغة البرمجة في الطرفين الأمامي والخلفي، مما يبسط التطوير والصيانة.')
    doc.add_paragraph('قاعدة البيانات المستخدمة هي SQLite لأنها تعمل في ملف محلي واحد ولا تحتاج إلى إعداد خادم قواعد بيانات كبير. هذا الخيار مناسب للمشروع التعليمي والتطبيقات الصغيرة حيث تكون سهولة التثبيت والتشغيل أهم من الأداء العالي جداً، كما تتيح حفظ البيانات بشكل موثوق دون ضرورة إدارة خوادم إضافية.')
    doc.add_section(WD_SECTION.NEW_PAGE)


# Add the section listing programming languages used in the project.
def add_languages(doc):
    doc.add_heading('اللغات المستخدمة', level=1)
    languages = [
        'JavaScript',
        'HTML',
        'CSS',
        'SQL',
    ]
    for lang in languages:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(lang)
    doc.add_section(WD_SECTION.NEW_PAGE)


# Add the database design section, including table schemas and ERD if available.
def add_database_design(doc):
    doc.add_heading('تصميم قاعدة البيانات', level=1)
    doc.add_paragraph('تستخدم قاعدة بيانات SQLite في ملف db/events.db. يعتمد تصميم البيانات على جداول مترابطة تدعم إدارة المستخدمين، الفعاليات، التذاكر، المدفوعات، والتعليقات.')
    tables = [
        {
            'name': 'events',
            'description': 'تخزين تفاصيل الفعاليات مثل العنوان، الوصف، الموقع، السعر، وحالة الفعالية.',
            'columns': [
                ('id', 'INTEGER PRIMARY KEY AUTOINCREMENT'),
                ('title', 'TEXT NOT NULL'),
                ('description', 'TEXT NOT NULL'),
                ('location', 'TEXT NOT NULL'),
                ('latitude', 'REAL'),
                ('longitude', 'REAL'),
                ('date', 'TEXT NOT NULL'),
                ('category', 'TEXT NOT NULL'),
                ('attendees', 'INTEGER NOT NULL DEFAULT 0'),
                ('archived', 'INTEGER NOT NULL DEFAULT 0'),
                ('image', 'TEXT'),
                ('isHybrid', 'INTEGER NOT NULL DEFAULT 0'),
                ('streamUrl', 'TEXT'),
                ('virtualPriceCents', 'INTEGER NOT NULL DEFAULT 1500'),
                ('budgetCents', 'INTEGER NOT NULL DEFAULT 0'),
                ('district', 'TEXT'),
            ]
        },
        {
            'name': 'users',
            'description': 'تخزين حسابات المستخدمين مع البريد الإلكتروني وكلمات المرور والحالة.',
            'columns': [
                ('id', 'INTEGER PRIMARY KEY AUTOINCREMENT'),
                ('username', 'TEXT UNIQUE NOT NULL'),
                ('email', 'TEXT UNIQUE NOT NULL'),
                ('password', 'TEXT NOT NULL'),
                ('role', "TEXT NOT NULL DEFAULT 'normal'"),
                ('isActivated', 'INTEGER NOT NULL DEFAULT 0'),
                ('activationToken', 'TEXT'),
                ('activationExpires', 'INTEGER'),
                ('resetPasswordToken', 'TEXT'),
                ('resetPasswordExpires', 'INTEGER'),
            ]
        },
        {
            'name': 'tickets',
            'description': 'تخزين بيانات التذاكر مع الربط بالفعلية والمستخدم، وحالة الدفع، ورمز QR.',
            'columns': [
                ('id', 'INTEGER PRIMARY KEY AUTOINCREMENT'),
                ('eventId', 'INTEGER NOT NULL'),
                ('userId', 'INTEGER'),
                ('username', 'TEXT'),
                ('ticketType', 'TEXT'),
                ('priceCents', 'INTEGER NOT NULL'),
                ('currency', "TEXT NOT NULL DEFAULT 'usd'"),
                ('paymentProvider', 'TEXT'),
                ('status', "TEXT NOT NULL DEFAULT 'pending'"),
                ('ticketCode', 'TEXT UNIQUE'),
                ('qrPath', 'TEXT'),
                ('pdfPath', 'TEXT'),
                ('isVirtual', 'INTEGER NOT NULL DEFAULT 0'),
                ('seatNumber', 'INTEGER'),
                ('seatCategory', 'TEXT'),
                ('streamUrl', 'TEXT'),
                ('createdAt', "INTEGER NOT NULL DEFAULT (strftime('%s','now'))"),
            ]
        },
        {
            'name': 'payments',
            'description': 'تسجيل تفاصيل المدفوعات المرتبطة بالتذاكر ومقدار المبلغ والحالة.',
            'columns': [
                ('id', 'INTEGER PRIMARY KEY AUTOINCREMENT'),
                ('ticketId', 'INTEGER'),
                ('provider', 'TEXT'),
                ('providerChargeId', 'TEXT'),
                ('amountCents', 'INTEGER'),
                ('currency', 'TEXT'),
                ('status', 'TEXT'),
                ('rawResponse', 'TEXT'),
                ('createdAt', "INTEGER NOT NULL DEFAULT (strftime('%s','now'))"),
            ]
        },
        {
            'name': 'comments',
            'description': 'تعليقات المستخدمين وتقييماتهم للفعاليات.',
            'columns': [
                ('id', 'INTEGER PRIMARY KEY AUTOINCREMENT'),
                ('eventId', 'INTEGER NOT NULL'),
                ('username', 'TEXT NOT NULL'),
                ('content', 'TEXT NOT NULL'),
                ('rating', 'INTEGER'),
                ('createdAt', "INTEGER NOT NULL DEFAULT (strftime('%s','now'))"),
            ]
        },
        {
            'name': 'attendance',
            'description': 'متابعة حضور المستخدمين للفعاليات ووضعهم الحالي.',
            'columns': [
                ('id', 'INTEGER PRIMARY KEY AUTOINCREMENT'),
                ('eventId', 'INTEGER NOT NULL'),
                ('userId', 'INTEGER NOT NULL'),
                ('username', 'TEXT NOT NULL'),
                ('status', 'TEXT NOT NULL'),
                ('updatedAt', "INTEGER NOT NULL DEFAULT (strftime('%s','now'))"),
            ]
        },
        {
            'name': 'event_media',
            'description': 'وسائط الفعالية مثل الصور والفيديو المرتبطة بكل حدث.',
            'columns': [
                ('id', 'INTEGER PRIMARY KEY AUTOINCREMENT'),
                ('eventId', 'INTEGER NOT NULL'),
                ('type', 'TEXT NOT NULL'),
                ('url', 'TEXT NOT NULL'),
                ('filename', 'TEXT'),
                ('position', 'INTEGER NOT NULL DEFAULT 0'),
            ]
        },
    ]
    for tbl in tables:
        doc.add_heading(tbl['name'], level=2)
        doc.add_paragraph(tbl['description'])
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = 'العمود'
        hdr[1].text = 'النوع والوصف'
        # style the table for better visual appearance
        try:
            style_table(table, col_widths=[2.5, 4.5], header_fill='D9D9D9')
        except Exception:
            pass
        for col_name, col_type in tbl['columns']:
            row = table.add_row().cells
            row[0].text = col_name
            row[1].text = col_type
        doc.add_paragraph('')
    erd_path = os.path.join(IMAGES_DIR, 'erd.png')
    if os.path.exists(erd_path):
        doc.add_paragraph('مخطط ERD يوضح العلاقات بين الجداول:')
        doc.add_picture(erd_path, width=Inches(6))
    doc.add_section(WD_SECTION.NEW_PAGE)


# Add the UI design section with descriptions of each front-end page.
def add_ui_design(doc):
    doc.add_heading('تصميم الواجهات', level=1)
    doc.add_paragraph('الواجهة تعتمد على صفحات HTML/CSS/JS لتقديم تجربة مستخدم تفاعلية ومتكيفة مع الهواتف. تُبنى الواجهات لتسهيل استعراض الفعاليات، إتمام الحجز، وعرض التذاكر. في ما يلي وصف تفصيلي لكل واجهة وعناصرها:')
    ui_items = [
        ('index.html', [
            'هدَف الصفحة: تمكين المستخدم من استعراض الفعاليات بسرعة وإيجاد ما يبحث عنه.',
            'مكونات رئيسية: شريط بحث، فلاتر (التاريخ، الفئة، المدينة)، شبكة بطاقات فعاليات، زر تفاصيل/حجز لكل بطاقة.',
            'سلوك تفاعلي: تحميل متزايد (infinite scroll) أو ترقيم صفحات، عرض حالة المقاعد المتبقية، روابط للمشاركة.',
            'إمكانيات إضافية: تحديد الموقع الجغرافي لعرض فعاليات قريبة.'
        ]),
        ('event.html', [
            'هدَف الصفحة: عرض تفاصيل الفعالية بالكامل للمستخدم قبل الحجز.',
            'مكونات رئيسية: عنوان، وصف طويل، صور/معرض، خريطة (embed)، جدول مواعيد، أسعار التذاكر وأنواعها، زر الحجز.',
            'سلوك تفاعلي: عرض تقييمات وتعليقات المستخدمين، اختيار عدد التذاكر، عرض السياسات.',
        ]),
        ('checkout.html', [
            'هدَف الصفحة: جمع بيانات الحجز والدفع بشكل آمن وميسّر.',
            'مكونات رئيسية: ملخص الحجز، نموذج بيانات حامل البطاقة، حقل كوبون خصم، زر تنفيذ الدفع.',
            'أمان: استخدام واجهة Stripe Elements أو تحويل آمن للبيانات لتفادي مرور بيانات البطاقة عبر الخادم.',
            'تجربة المستخدم: عرض حالة الطلب أثناء المعالجة، رسائل خطأ واضحة.'
        ]),
        ('profile.html', [
            'هدَف الصفحة: إدارة بيانات المستخدم وحجوزاته السابقة.',
            'مكونات رئيسية: معلومات الحساب، قائمة الحجوزات مع حالة كل حجز، روابط لتحميل التذاكر (PDF)، خيار إلغاء الحجز عند الإمكان.',
            'خصوصية: إظهار/إخفاء بيانات حساسة بناء على صلاحيات المستخدم.'
        ])
    ]
    for filename, bullets in ui_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(filename + ':')
        for b in bullets:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run(b)
    ui_images = create_ui_images()
    for img_path in ui_images:
        doc.add_paragraph(f'صورة العرض: {os.path.basename(img_path)}')
        doc.add_picture(img_path, width=Inches(6))
    doc.add_section(WD_SECTION.NEW_PAGE)


# Add the use case section that describes how different actors interact with the system.
def add_use_case_section(doc):
    doc.add_heading('مخطط حالات الاستخدام', level=1)
    doc.add_paragraph('يوضح هذا المخطط السيناريوهات الأساسية التي يمكن للمستخدمين القيام بها في النظام. لكل حالة استخدام نعرض الممثل (Actor)، الشروط المسبقة، التدفق الرئيسي، والتفرعات المحتملة.')
    use_cases = [
        {
            'title': 'عرض الفعاليات',
            'actor': 'زائر / مستخدم',
            'pre': 'الاتصال بالإنترنت، وجود فعاليات منشورة في النظام',
            'page': 'index.html',
            'main': [
                'يفتح المستخدم الصفحة الرئيسية لعرض الفعاليات المتاحة.',
                'يعرض النظام قائمة الفعاليات مرتبة أو بحسب عوامل التصفية.',
                'يستخدم المستخدم شريط البحث والفلاتر للعثور على فعالية مناسبة.',
                'يضغط المستخدم على بطاقة الفعالية لعرض التفاصيل.'
            ],
            'alt': ['فلتر النتائج لا تتطابق مع أي فعالية → يعرض النظام رسالة "لا توجد فعاليات" مع اقتراح إزالة الفلاتر.']
        },
        {
            'title': 'حجز تذكرة',
            'actor': 'مستخدم مسجل',
            'pre': 'المستخدم مسجل وفعّال، والفعالية بها تذاكر متاحة',
            'page': 'event.html',
            'main': [
                'يختار المستخدم نوع وكمية التذاكر من صفحة تفاصيل الفعالية.',
                'يعرض النظام ملخص الحجز والسعر النهائي.',
                'يضغط المستخدم على زر الحجز للانتقال إلى صفحة الدفع.'
            ],
            'alt': ['المقاعد غير متوفرة أثناء الحجز → يظهر خطأ ويعود المستخدم لاختيار كمية أقل أو فعالية أخرى.']
        },
        {
            'title': 'الدفع عبر Stripe',
            'actor': 'مستخدم / النظام (بوابة الدفع)',
            'pre': 'بيانات البطاقة صالحة، والاتصال بـ Stripe متاح',
            'page': 'checkout.html',
            'main': [
                'يُحمل نموذج الدفع عبر Stripe Elements في صفحة الدفع الآمنة.',
                'يُدخل المستخدم بيانات البطاقة ويضغط على تأكيد الدفع.',
                'يتحقق النظام من نتيجة الدفع ويحدث حالة الحجز إلى مدفوع أو فشل.'
            ],
            'alt': ['فشل الدفع → يعرض المستخدم رسالة خطأ ويتيح له إعادة المحاولة أو اختيار وسيلة دفع أخرى.']
        },
        {
            'title': 'توليد تذكرة QR',
            'actor': 'النظام',
            'pre': 'تم تأكيد الدفع وتم حفظ الحجز',
            'page': 'profile.html',
            'main': [
                'يتحقق النظام من نجاح الدفع وحالة الحجز.',
                'ينشئ رمز QR فريد مرتبطاً بتفاصيل التذكرة.',
                'يحفظ النظام ملف PDF للتذكرة ويرسل رابط التنزيل للمستخدم.'
            ],
            'alt': []
        }
    ]
    for uc in use_cases:
        doc.add_heading(uc['title'], level=2)
        doc.add_paragraph(f"الممثل: {uc['actor']}")
        doc.add_paragraph(f"الواجهة المرتبطة: {uc['page']}")
        doc.add_paragraph(f"الشروط المسبقة: {uc['pre']}")
        doc.add_paragraph('التدفق الرئيسي:')
        for step in uc['main']:
            p = doc.add_paragraph(style='List Number')
            p.add_run(step)
        if uc['alt']:
            doc.add_paragraph('التفرعات والحالات البديلة:')
            for alt in uc['alt']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(alt)
    use_case_path = os.path.join(IMAGES_DIR, 'use_case.png')
    if os.path.exists(use_case_path):
        doc.add_paragraph('المخطط العام لحالات الاستخدام:')
        doc.add_picture(use_case_path, width=Inches(6))
    doc.add_section(WD_SECTION.NEW_PAGE)


# Add a section that explains key workflows and algorithms used in the project.
def add_algorithms(doc):
    doc.add_heading('الخوارزميات والعمليات', level=1)
    doc.add_paragraph('يتضمن المشروع عدة عمليات أساسية مع خوارزميات مبسطة لكل عملية لتوضيح التسلسل المنطقي والتنفيذ.')
    # Booking flow pseudocode
    doc.add_heading('تدفق الحجز (مبسّط)', level=2)
    doc.add_paragraph('وصف مختصر: التحقق من توافر التذاكر، حجز مقعد مؤقت، معالجة الدفع، تأكيد الحجز وتوليد التذكرة.')
    booking_steps = [
        '1. استعلام عن توافر التذاكر حسب نوع السعر والكمية المطلوبة.',
        '2. حجز مقعد مؤقت (reserve) لتفادي البيع المتزامن.',
        '3. تهيئة جلسة دفع وإعادة توجيه المستخدم لواجهة الدفع الآمنة.',
        '4. تفعيل webhook من مزود الدفع للتحقق من نجاح العملية.',
        '5. عند النجاح: تأكيد الحجز، توليد رمز QR، وإرسال التذكرة بالبريد الإلكتروني.'
    ]
    for s in booking_steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(s)

    doc.add_heading('خوارزمية معالجة الدفع (موجز)', level=2)
    pay_steps = [
        'توليد طلب دفع (PaymentIntent) في Stripe مع مبلغ بالملّيمات.',
        'عرض واجهة الدفع الآمنة (Stripe Elements) وجمع بيانات البطاقة.',
        'تنفيذ الدفع واستقبال نتيجة العملية.',
        'تخزين نتيجة الدفع وربطها بسجل التذكرة (status = paid/failed).'
    ]
    for s in pay_steps:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(s)

    doc.add_heading('توليد تذاكر QR', level=2)
    qr_steps = [
        'إنشاء رمز فريد (UUID) مرتبط بمعرف الحجز.',
        'تشفير أو توقيع بيانات التذكرة عند الحاجة لمنع التزييف.',
        'توليد صورة QR وحفظها مع ملف PDF للتذكرة.',
        'عرض رابط التنزيل وإرساله في رسالة تأكيد.'
    ]
    for s in qr_steps:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(s)
    doc.add_section(WD_SECTION.NEW_PAGE)


# Add the test cases section with example scenarios and expected results.
def add_testing(doc):
    doc.add_heading('الاختبارات', level=1)
    doc.add_paragraph('يتضمن هذا القسم حالات اختبارية مفصّلة تشمل خطوات التنفيذ، المدخلات، والنتيجة المتوقعة لكل حالة.')
    tests = [
        ('تسجيل مستخدم جديد', 'نموذج التسجيل: اسم، بريد إلكتروني صالح، كلمة مرور', 'إنشاء سجل مستخدم في جدول users، إرسال إيميل تفعيل، عرض رسالة نجاح'),
        ('تأكيد الحساب', 'زيارة رابط التفعيل من البريد', 'تحديث الحقل isActivated = 1، إعادة توجيه لصفحة تسجيل الدخول'),
        ('تسجيل الدخول', 'إدخال إيميل وكلمة مرور صحيحة', 'تسجيل الدخول بنجاح، جلسة مستخدم صالحة'),
        ('حجز تذكرة ودفع ناجح', 'اختيار فعالية، تحديد تذاكر، إجراء دفع باستخدام بطاقة اختبار Stripe', 'حالة الحجز = paid، إنشاء سجل في tickets، توليد PDF مع QR، إرسال بريد تأكيد'),
        ('حجز تذكرة وفشل الدفع', 'دفع يفشل أو يتم إلغاؤه', 'حالة الحجز = pending/failed، عدم إنشاء تذكرة، رسالة خطأ للمستخدم'),
        ('عرض فعاليات بفلتر', 'تطبيق فلتر التاريخ أو المدينة', 'إظهار قائمة منقّحة تطابق شروط البحث، أداء استعلام مناسب'),
    ]
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = 'حالة الاختبار'
    hdr[1].text = 'المدخلات'
    hdr[2].text = 'النتيجة المتوقعة'
    try:
        style_table(table, col_widths=[2.5, 4.0, 4.0], header_fill='D9D9D9')
    except Exception:
        pass
    for name, inp, expected in tests:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = inp
        row[2].text = expected

    doc.add_paragraph('الصور التالية توضح واجهات المستخدم المرتبطة ببعض حالات الاختبار وضمان صحة التدفق المرئي:')
    ui_images = {
        'ui_index.png': 'صفحة العرض الرئيسية لحالة اختبار عرض الفعاليات بالفلتر. يظهر شريط البحث وقائمة الفعاليات.',
        'ui_event.png': 'صفحة تفاصيل الفعالية لحالة اختبار حجز التذكرة. تعرض المعلومات الأساسية وخيارات التذكرة.',
        'ui_checkout.png': 'صفحة الدفع لحالة اختبار الدفع الناجح أو الفاشل عبر Stripe. توضح نموذج الدفع الآمن.',
        'ui_profile.png': 'صفحة حساب المستخدم لحالة اختبار متابعة الحجوزات، عرض التذاكر، واستعراض الحالة.'
    }
    for img_name, caption in ui_images.items():
        img_path = os.path.join(IMAGES_DIR, img_name)
        if os.path.exists(img_path):
            doc.add_paragraph(f'صورة الواجهة: {img_name}')
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph(caption)
    doc.add_section(WD_SECTION.NEW_PAGE)


# Add the conclusion and recommendations section of the report.
def add_conclusions(doc):
    doc.add_heading('الاستنتاجات والتوصيات', level=1)
    doc.add_paragraph('هذا المشروع يمثل خطوة عملية نحو تطوير نظام فعاليات متكامل لإدارة الحجوزات، التذاكر، ودفع التكاليف عبر الإنترنت. يمكن تطويره لاحقاً ليدعم أكثر من جهة تنظيم ويفتح صلاحيات لإدارة متعددة. كما ينصح بتحسين واجهات الاستخدام لتكون أكثر تجاوباً مع الهواتف الذكية. ')
    doc.add_section(WD_SECTION.NEW_PAGE)


# Add an appendix listing the main project files included in the report.
def add_appendix(doc, files_info):
    doc.add_heading('ملحق الملفات', level=1)
    doc.add_paragraph('قائمة الملفات الأساسية في المشروع مع معلومات موجزة عن كل منها:')
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = 'اسم الملف'
    hdr[1].text = 'النوع'
    hdr[2].text = 'ملاحظة'
    try:
        style_table(table, col_widths=[5.0, 1.0, 2.0], header_fill='D9D9D9')
    except Exception:
        pass
    for _, rel in files_info:
        row = table.add_row().cells
        row[0].text = rel
        row[1].text = os.path.splitext(rel)[1].lstrip('.')
        row[2].text = 'ملف مشروع'.ljust(0)
    doc.add_section(WD_SECTION.NEW_PAGE)


# Generate the Word (.docx) version of the report from collected data and images.
def make_docx(files_info, package_info):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    def apply_document_style(d):
        # Set Normal style
        try:
            normal = d.styles['Normal']
            if ARABIC_FONT_PATH:
                normal.font.name = os.path.splitext(os.path.basename(ARABIC_FONT_PATH))[0]
            else:
                normal.font.name = 'Times New Roman'
            normal.font.size = Pt(13)
        except Exception:
            pass

        # Heading 1 - centered title
        try:
            h1 = d.styles['Heading 1']
            h1.font.size = Pt(20)
            h1.font.bold = True
            if ARABIC_FONT_PATH:
                h1.font.name = os.path.splitext(os.path.basename(ARABIC_FONT_PATH))[0]
            h1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception:
            pass

        # Heading 2 - section titles (right aligned)
        try:
            h2 = d.styles['Heading 2']
            h2.font.size = Pt(16)
            h2.font.bold = True
            if ARABIC_FONT_PATH:
                h2.font.name = os.path.splitext(os.path.basename(ARABIC_FONT_PATH))[0]
            h2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        except Exception:
            pass

        # Heading 3 - subsection
        try:
            h3 = d.styles['Heading 3']
            h3.font.size = Pt(14)
            h3.font.bold = True
            if ARABIC_FONT_PATH:
                h3.font.name = os.path.splitext(os.path.basename(ARABIC_FONT_PATH))[0]
            h3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        except Exception:
            pass

    apply_document_style(doc)

    # add header and footer with project title and page numbers
    def add_header_footer(d):
        for section in d.sections:
            # header with project title centered
            try:
                hdr = section.header
                if hdr.paragraphs:
                    p = hdr.paragraphs[0]
                else:
                    p = hdr.add_paragraph()
                p.text = PROJECT_TITLE
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except Exception:
                pass
            # footer with centered page number
            try:
                ftr = section.footer
                if ftr.paragraphs:
                    p = ftr.paragraphs[0]
                else:
                    p = ftr.add_paragraph()
                p.text = ''
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # insert a PAGE field
                fld = OxmlElement('w:fldSimple')
                fld.set(qn('w:instr'), 'PAGE')
                p._p.append(fld)
            except Exception:
                pass

    add_header_footer(doc)

    add_cover_page(doc)
    add_contents_page(doc)
    add_dedication(doc)
    add_introduction(doc)
    add_objectives(doc)
    add_advantages(doc)
    add_technologies(doc, package_info)
    add_code_explanation(doc)
    add_languages(doc)
    add_database_design(doc)
    add_ui_design(doc)
    add_use_case_section(doc)
    add_algorithms(doc)
    add_testing(doc)
    add_conclusions(doc)
    add_appendix(doc, files_info)
    # re-apply header/footer after sections have been added so each section gets header/footer
    add_header_footer(doc)
    # Save to a temporary file first to avoid write errors when the target is locked.
    tmp = OUTPUT_DOCX + '.tmp'
    try:
        doc.save(tmp)
    except Exception:
        # final fallback: try to save directly (may raise PermissionError)
        doc.save(OUTPUT_DOCX)
        return OUTPUT_DOCX
    # Try to replace the target atomically; if replacement fails (e.g., file is open), keep tmp and return its path.
    try:
        os.replace(tmp, OUTPUT_DOCX)
        return OUTPUT_DOCX
    except Exception:
        return tmp


# Generate the PDF version of the report using fpdf for simple layout.
def make_pdf(package_info):
    from fpdf import FPDF
    pdf = FPDF(format='A4')
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    arial_path = r'C:\Windows\Fonts\arial.ttf'
    if os.path.exists(arial_path):
        pdf.add_font('Arial', '', arial_path, uni=True)
        pdf.add_font('Arial', 'B', arial_path, uni=True)
    else:
        pdf.add_font('Arial', '', '', uni=True)
        pdf.add_font('Arial', 'B', '', uni=True)
    pdf.set_font('Arial', '', 14)
    # Helper to add a titled section to the PDF with optional text and bullet list.
    def add_section(title, text_lines=None, bullets=None):
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, title, ln=True)
        pdf.ln(2)
        pdf.set_font('Arial', '', 12)
        if text_lines:
            for line in text_lines:
                pdf.multi_cell(0, 8, line)
            pdf.ln(2)
        if bullets:
            for bullet in bullets:
                pdf.multi_cell(0, 8, f'• {bullet}')
            pdf.ln(2)

    title_lines = [
        'وزارة التعليم العالي',
        'جامعة دمشق',
        INSTITUTE_TITLE,
        DEPARTMENT_TITLE,
        PROJECT_TITLE,
    ]
    pdf.set_font('Arial', 'B', 20)
    for line in title_lines:
        pdf.cell(0, 12, line, ln=True, align='C')
    pdf.ln(8)
    pdf.set_font('Arial', '', 14)
    pdf.cell(0, 10, f'أعدّه الطلاب التالية أسماؤهم:', ln=True, align='C')
    for name in TEAM_MEMBERS:
        pdf.cell(0, 8, name, ln=True, align='C')
    pdf.ln(10)
    pdf.cell(0, 10, f'العام الدراسي {ACADEMIC_YEAR}', ln=True, align='C')
    pdf.add_page()
    add_section('المقدمة', text_lines=[ARABIC_INTRO])
    add_section('الهدف من المشروع', bullets=[
        'إنشاء نظام إلكتروني لإدارة فعاليات محافظة درعا بطريقة متكاملة.',
        'تيسير حجز التذاكر ومتابعة الحجوزات الكترونياً.',
        'إدارة بيانات الفعالية والمستخدمين بأمان وبدون تعقيد.',
        'توليد تذاكر QR وربط الدفع الالكتروني.',
    ])
    add_section('مزايا المشروع', bullets=[
        'زيادة الكفاءة وتقليل الاعتماد على الأوراق.',
        'تحسين تجربة المستخدم وسهولة الحجز.',
        'توفير أدوات إدارية لمتابعة المبيعات.',
        'دعم الدفع الإلكتروني عبر Stripe.',
    ])
    pdf.add_page()
    add_section('التقنيات المستخدمة', bullets=[
        'Node.js', 'Express', 'SQLite', 'HTML', 'CSS', 'JavaScript', 'Stripe', 'PDFKit', 'QRCode', 'Nodemailer', 'Multer',
    ])
    add_section('شرح الكود', text_lines=[
        'الواجهة الأمامية (Front-end) تعتمد على HTML وCSS وJavaScript لتقديم تجربة مستخدم تفاعلية وسهلة الاستخدام. يتم استخدام صفحات HTML لهيكلة المحتوى، وCSS لتصميم الشبكة والألوان، وJavaScript لإضافة التفاعلية مثل عرض الفعاليات، حقول البحث، وانتقال المستخدم بين الصفحات دون إعادة تحميل كاملة.',
        'الجهة الخلفية (Back-end) تستخدم Node.js مع إطار Express لبناء خادم ويب بسيط يعالج طلبات HTTP، يستعلم عن البيانات من قاعدة البيانات، ويُدير جلسات المستخدم وتكامل الدفع. تم اختيار Node.js لأنه يُسهّل كتابة كود جافاسكربت واحد لكل من الواجهة الأمامية والخلفية، مما يُقلل التعقيد ويوفر تطبيقاً أسرع في التطوير.',
        'قاعدة البيانات تستخدم SQLite بسبب سهولة إعدادها وعدم الحاجة لتشغيل خادم قواعد بيانات منفصل. هذا الخيار مناسب للمشروع التعليمي والتطبيقات الصغيرة حيث يُمكن حفظ البيانات محلياً في ملف واحد، مع أداء جيد لقراءة وكتابة البيانات المتوسطة.',
    ])
    add_section('اللغات المستخدمة', bullets=['JavaScript', 'HTML', 'CSS', 'SQL'])
    pdf.add_page()
    add_section('تصميم قاعدة البيانات', text_lines=['القاعدة تستخدم جدولاً لتنظيم المستخدمين، الفعاليات، الحجوزات، والتذاكر.'])
    erd_path = os.path.join(IMAGES_DIR, 'erd.png')
    if os.path.exists(erd_path):
        pdf.image(erd_path, w=170)
    pdf.add_page()
    add_section('تصميم الواجهات', bullets=[
        'index.html لعرض الفعاليات.',
        'event.html لتفاصيل الفعالية.',
        'checkout.html لإتمام الحجز والدفع.',
        'profile.html لإدارة حساب المستخدم.',
    ])
    ui_samples = ['ui_index.png', 'ui_event.png', 'ui_checkout.png', 'ui_profile.png']
    for sample in ui_samples:
        img_path = os.path.join(IMAGES_DIR, sample)
        if os.path.exists(img_path):
            pdf.image(img_path, w=170)
            pdf.ln(5)
    add_section('الخوارزميات والعمليات', bullets=[
        'التحقق من هوية المستخدم.',
        'إدارة الحجوزات وتحديث حالة المقاعد.',
        'توليد تذاكر PDF مع رمز QR.',
        'تكامل الدفع عبر Stripe.',
    ])
    pdf.add_page()
    use_case_path = os.path.join(IMAGES_DIR, 'use_case.png')
    if os.path.exists(use_case_path):
        pdf.image(use_case_path, w=170)
        pdf.ln(10)
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, 'مخطط حالات الاستخدام', ln=True)
        pdf.ln(4)
        pdf.set_font('Arial', '', 12)
        pdf.multi_cell(0, 8, 'يوضح هذا المخطط السيناريوهات الرئيسية مثل عرض الفعالية، حجز التذكرة، الدفع، وإدارة الحساب.')
    pdf.add_page()
    add_section('الاختبارات', bullets=[
        'اختبار تسجيل المستخدم.',
        'اختبار الحجز وتوليد التذاكر.',
        'اختبار الدفع والتكامل مع Stripe.',
        'اختبار عرض وتحديث الفعاليات.',
    ])
    add_section('الاستنتاجات والتوصيات', text_lines=['المشروع يوفر بنية متينة لإدارة الفعاليات والمبيعات، كما يمكن تطويره لإضافة تقسيمات إدارية متعددة وتحسين تجربة المستخدم على الأجهزة المحمولة.'])
    # write to temp then try to replace final pdf
    tmp_pdf = OUTPUT_PDF + '.tmp'
    try:
        pdf.output(tmp_pdf)
    except Exception:
        pdf.output(OUTPUT_PDF)
        return OUTPUT_PDF
    try:
        os.replace(tmp_pdf, OUTPUT_PDF)
        return OUTPUT_PDF
    except Exception:
        return tmp_pdf


# Create a ZIP archive of the generated report files.
def make_zip(paths, zip_path):
    tmp_zip = zip_path + '.tmp'
    try:
        if os.path.exists(tmp_zip):
            os.remove(tmp_zip)
    except Exception:
        pass
    with zipfile.ZipFile(tmp_zip, 'w', compression=zipfile.ZIP_DEFLATED) as z:
        for p in paths:
            if os.path.exists(p):
                z.write(p, arcname=os.path.basename(p))
    try:
        os.replace(tmp_zip, zip_path)
        return zip_path
    except Exception:
        return tmp_zip


# Main entry point: build images, the DOCX report, the PDF report, and package them if needed.
def main():
    os.makedirs(IMAGES_DIR, exist_ok=True)
    erd_path = os.path.join(IMAGES_DIR, 'erd.png')
    if not os.path.exists(erd_path):
        created = create_sample_erd(erd_path)
        if created:
            print('تم إنشاء صورة ERD توضيحية في images/erd.png')

    package_info = read_package_json()
    files = collect_files(WORKSPACE_ROOT)
    files_info = []
    for full, rel in files:
        info = summarize_file(full)
        info['relpath'] = rel
        files_info.append((full, rel))

    generated_ui = create_ui_images()
    real_screens = collect_real_screenshots()
    if real_screens:
        print(f'تم نسخ {len(real_screens)} لقطة شاشة حقيقية من screenshots/ إلى images/')
    # prefer real screenshots in images folder; also generate use case diagram
    create_use_case_diagram(os.path.join(IMAGES_DIR, 'use_case.png'))

    docx_path = make_docx(files_info, package_info)
    print(f'تم إنشاء ملف Word: {docx_path}')

    pdf_path = make_pdf(package_info)
    print(f'تم إنشاء ملف PDF: {pdf_path}')

    zip_path = make_zip([docx_path, pdf_path], OUTPUT_ZIP)
    print(f'تم إنشاء ملف ZIP: {zip_path}')


if __name__ == '__main__':
    main()
