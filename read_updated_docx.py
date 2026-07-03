"""Script to read التقرير-1-UPDATED.docx"""
from docx import Document

DOC = 'التقرير-1-UPDATED.docx'
try:
    doc = Document(DOC)
except Exception as e:
    print('ERROR_OPEN', e)
    raise SystemExit(1)

print('PARAGRAPHS_COUNT:', len(doc.paragraphs))
print('\n--- All paragraphs ---\n')
for i, p in enumerate(doc.paragraphs, 1):
    if p.text.strip():  # Only print non-empty paragraphs
        print(f'[{i}] {p.text}')
