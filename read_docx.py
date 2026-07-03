"""Utility script to inspect the generated Word report.
This script opens the output DOCX file, prints the paragraph count,
shows the first 40 paragraphs, and lists embedded media files.
"""

from docx import Document
from zipfile import ZipFile

DOC = 'التقرير-1.docx'
try:
    doc = Document(DOC)
except Exception as e:
    print('ERROR_OPEN', e)
    raise SystemExit(1)

print('PARAGRAPHS_COUNT:', len(doc.paragraphs))
print('\n--- First 40 paragraphs ---\n')
for i, p in enumerate(doc.paragraphs[:40], 1):
    print(f'[{i}]', p.text)

print('\n--- Embedded images in word/media/ ---')
try:
    with ZipFile(DOC) as z:
        imgs = [n for n in z.namelist() if n.startswith('word/media/')]
        print('IMAGES_COUNT:', len(imgs))
        for im in imgs:
            print('-', im)
except Exception as e:
    print('ERROR_LIST_IMAGES', e)
